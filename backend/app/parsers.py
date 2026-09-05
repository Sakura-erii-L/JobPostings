from __future__ import annotations

import ast
import csv
import hashlib
import io
import ipaddress
import json
import mimetypes
import re
import socket
from datetime import datetime, timedelta, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urljoin, urlparse
from zoneinfo import ZoneInfo

import httpx


SOURCE_TIMEZONE = ZoneInfo("Asia/Shanghai")
MESSAGE_TIME_FIELDS = (
    "datetime",
    "dateTime",
    "message_datetime",
    "messageDatetime",
    "sent_at",
    "sentAt",
    "message_time",
    "messageTime",
    "time",
    "timestamp",
)
MESSAGE_TIME_FORMATS = (
    "%Y/%m/%d %H:%M:%S",
    "%Y/%m/%d %H:%M",
    "%Y-%m-%d %H:%M:%S",
    "%Y-%m-%d %H:%M",
    "%Y.%m.%d %H:%M:%S",
    "%Y.%m.%d %H:%M",
    "%Y年%m月%d日 %H:%M:%S",
    "%Y年%m月%d日 %H:%M",
    "%Y/%m/%d",
    "%Y-%m-%d",
)
EVENT_TIME_FORMATS = (
    "%Y/%m/%d %H:%M:%S",
    "%Y/%m/%d %H:%M",
    "%Y/%m/%d",
    "%Y-%m-%d %H:%M:%S",
    "%Y-%m-%d %H:%M",
    "%Y-%m-%d",
    "%Y.%m.%d %H:%M:%S",
    "%Y.%m.%d %H:%M",
    "%Y.%m.%d",
    "%Y年%m月%d日 %H时%M分%S秒",
    "%Y年%m月%d日 %H时%M分",
    "%Y年%m月%d日 %H:%M:%S",
    "%Y年%m月%d日 %H:%M",
    "%Y年%m月%d日",
)


def _parse_message_time_value(value: Any) -> str | None:
    """Normalize an explicit human-readable message time; never parse numeric epochs."""
    if not isinstance(value, str):
        return None
    text = value.strip()
    if not text or re.fullmatch(r"\d+(?:\.\d+)?", text):
        return None
    candidate = text[:-1] + "+00:00" if text.endswith(("Z", "z")) else text
    try:
        parsed = datetime.fromisoformat(candidate)
    except ValueError:
        parsed = None
        for time_format in MESSAGE_TIME_FORMATS:
            try:
                parsed = datetime.strptime(text, time_format)
                break
            except ValueError:
                continue
        if parsed is None:
            return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=SOURCE_TIMEZONE)
    return parsed.astimezone(timezone.utc).isoformat(timespec="seconds")


def parse_message_time(message: dict[str, Any]) -> str | None:
    """Return the source chat time from an explicit date field, or None when unknown."""
    for field in MESSAGE_TIME_FIELDS:
        parsed = _parse_message_time_value(message.get(field))
        if parsed:
            return parsed
    return None


def _parse_explicit_datetime(value: Any, timezone_name: str = "Asia/Shanghai") -> datetime | None:
    """Parse a human date while never treating numeric metadata as an epoch."""
    if value is None or isinstance(value, (int, float)):
        return None
    text = str(value).strip()
    if not text or re.fullmatch(r"\d+(?:\.\d+)?", text):
        return None
    candidate = text[:-1] + "+00:00" if text.endswith(("Z", "z")) else text
    try:
        parsed = datetime.fromisoformat(candidate)
    except ValueError:
        parsed = None
        for time_format in EVENT_TIME_FORMATS:
            try:
                parsed = datetime.strptime(text, time_format)
                break
            except ValueError:
                continue
    if parsed is None:
        return None
    if parsed.tzinfo is None:
        try:
            parsed = parsed.replace(tzinfo=ZoneInfo(timezone_name))
        except Exception:
            parsed = parsed.replace(tzinfo=SOURCE_TIMEZONE)
    return parsed.astimezone(timezone.utc)


def _event_zone(timezone_name: str) -> ZoneInfo:
    try:
        return ZoneInfo(timezone_name)
    except Exception:
        return SOURCE_TIMEZONE


def _clock_parts(text: str) -> tuple[int, int, int] | None:
    match = re.search(
        r"(?P<ampm>上午|下午|早上|晚上|夜间)?\s*(?P<hour>\d{1,2})"
        r"\s*(?:(?:[:：时点])\s*(?P<minute>\d{1,2}))?\s*(?:分)?",
        text,
    )
    if not match:
        return None
    hour = int(match.group("hour"))
    minute = int(match.group("minute") or 0)
    ampm = match.group("ampm") or ""
    if ampm in {"下午", "晚上", "夜间"} and hour < 12:
        hour += 12
    if ampm in {"上午", "早上"} and hour == 12:
        hour = 0
    if hour > 23 or minute > 59:
        return None
    return hour, minute, 0


def _parse_relative_datetime(value: Any, timezone_name: str, reference_at: str | None) -> datetime | None:
    if not reference_at:
        return None
    text = str(value or "").strip()
    match = re.search(r"今天|今日|当天|本日|明天|明日|次日|翌日|后天|后日|大后天", text)
    if not match:
        return None
    reference = _parse_explicit_datetime(reference_at, timezone_name)
    if reference is None:
        return None
    offsets = {
        "今天": 0, "今日": 0, "当天": 0, "本日": 0,
        "明天": 1, "明日": 1, "次日": 1, "翌日": 1,
        "后天": 2, "后日": 2, "大后天": 3,
    }
    zone = _event_zone(timezone_name)
    local_reference = reference.astimezone(zone)
    target_date = local_reference.date() + timedelta(days=offsets[match.group(0)])
    clock = _clock_parts(text[match.end():]) or (0, 0, 0)
    return datetime(*target_date.timetuple()[:3], *clock, tzinfo=zone).astimezone(timezone.utc)


def _parse_month_day_datetime(value: Any, timezone_name: str, reference_at: str | None) -> datetime | None:
    text = str(value or "").strip()
    match = re.search(
        r"(?<!\d)(?P<month>\d{1,2})\s*(?:月|[./-])\s*(?P<day>\d{1,2})\s*(?:日|号)?",
        text,
    )
    if not match:
        return None
    reference = _parse_explicit_datetime(reference_at, timezone_name) if reference_at else None
    zone = _event_zone(timezone_name)
    local_reference = (reference or datetime.now(timezone.utc)).astimezone(zone)
    clock = _clock_parts(text[match.end():]) or (0, 0, 0)
    try:
        return datetime(local_reference.year, int(match.group("month")), int(match.group("day")), *clock, tzinfo=zone).astimezone(timezone.utc)
    except ValueError:
        return None


def _parse_event_datetime(value: Any, timezone_name: str, reference_at: str | None) -> datetime | None:
    return (
        _parse_explicit_datetime(value, timezone_name)
        or _parse_relative_datetime(value, timezone_name, reference_at)
        or _parse_month_day_datetime(value, timezone_name, reference_at)
    )


def is_plausible_event_datetime(value: Any, reference_at: str | None = None, *, now: datetime | None = None) -> bool:
    """Reject dates that are implausibly far from the source message.

    Recruitment events can be in the past or future, but an event date many
    years away from the message is generally a metadata/model parsing error.
    The source message date is only a validation reference; it is never used
    as the event date.
    """
    parsed = _parse_explicit_datetime(value)
    if parsed is None:
        return False
    current = now or datetime.now(timezone.utc)
    if parsed.year < 1970 or parsed > current + timedelta(days=366 * 5):
        return False
    if reference_at:
        reference = _parse_explicit_datetime(reference_at)
        if reference is not None and abs((parsed - reference).total_seconds()) > 366 * 5 * 24 * 60 * 60:
            return False
    elif parsed < current - timedelta(days=366 * 10):
        # Without the source-message timestamp, very old dates are still
        # almost certainly metadata/model parsing errors for recruitment data.
        return False
    return True


def normalize_event_datetime(value: Any, timezone_name: str = "Asia/Shanghai", reference_at: str | None = None) -> str | None:
    """Return a trusted event time in UTC ISO format, or None if implausible."""
    parsed = _parse_event_datetime(value, timezone_name, reference_at)
    if parsed is None:
        return None
    if not is_plausible_event_datetime(parsed.isoformat(), reference_at):
        return None
    return parsed.isoformat(timespec="seconds")


def extract_event_datetime_candidates(text: str, timezone_name: str = "Asia/Shanghai", reference_at: str | None = None) -> list[str]:
    """Extract explicit dates from source text for repairing bad model dates."""
    if not text:
        return []
    candidates: list[str] = []
    full_date = re.compile(
        r"(?P<year>20\d{2})\s*(?:年|[./-])\s*(?P<month>\d{1,2})\s*(?:月|[./-])\s*(?P<day>\d{1,2})\s*"
        r"(?:(?P<ampm>上午|下午|早上|晚上|夜间)\s*)?(?P<hour>\d{1,2})?\s*(?:[:：时]\s*(?P<minute>\d{1,2}))?\s*(?:分\s*(?P<second>\d{1,2})\s*秒?)?"
    )
    month_day = re.compile(
        r"(?<!\d)(?P<month>\d{1,2})\s*月\s*(?P<day>\d{1,2})\s*(?:日|号)\s*"
        r"(?:(?P<ampm>上午|下午|早上|晚上|夜间)\s*)?(?P<hour>\d{1,2})?\s*(?:[:：时]\s*(?P<minute>\d{1,2}))?"
    )
    reference = _parse_explicit_datetime(reference_at) if reference_at else None
    reference_year = reference.year if reference else datetime.now(SOURCE_TIMEZONE).year

    def add_match(match: re.Match[str], year: int) -> None:
        try:
            hour = int(match.groupdict().get("hour") or 0)
            minute = int(match.groupdict().get("minute") or 0)
            second = int(match.groupdict().get("second") or 0)
            ampm = match.groupdict().get("ampm") or ""
            if ampm in {"下午", "晚上", "夜间"} and hour < 12:
                hour += 12
            candidate = datetime(
                year,
                int(match.group("month")),
                int(match.group("day")),
                hour,
                minute,
                second,
            )
        except (TypeError, ValueError):
            return
        normalized = normalize_event_datetime(candidate.isoformat(), timezone_name, reference_at)
        if normalized and normalized not in candidates:
            candidates.append(normalized)

    for match in full_date.finditer(text):
        add_match(match, int(match.group("year")))
    for match in month_day.finditer(text):
        add_match(match, reference_year)
    for match in re.finditer(r"(?:今天|今日|当天|本日|明天|明日|次日|翌日|后天|后日|大后天)(?:\s*(?:上午|下午|早上|晚上|夜间)?\s*\d{1,2}(?:(?:[:：时点])\s*\d{1,2})?\s*(?:分)?)?", text):
        normalized = normalize_event_datetime(match.group(0), timezone_name, reference_at)
        if normalized and normalized not in candidates:
            candidates.append(normalized)
    return candidates


def _catalog_section(text: str, heading_pattern: str) -> str:
    lines = text.replace("\r", "").split("\n")
    heading = re.compile(heading_pattern, re.IGNORECASE)
    numbered_heading = re.compile(r"^\s*(?:[一二三四五六七八九十百]+|\d+)[、.．:：]\s*\S+")
    start = None
    for index, line in enumerate(lines):
        if heading.search(line.strip()):
            start = index + 1
            break
    if start is None:
        return ""
    end = len(lines)
    for index in range(start, len(lines)):
        if numbered_heading.match(lines[index].strip()):
            end = index
            break
    return "\n".join(lines[start:end]).strip()


def _split_top_level(text: str, separators: str) -> list[str]:
    values: list[str] = []
    current: list[str] = []
    depth = 0
    for character in text:
        if character in "（(【[":
            depth += 1
        elif character in "）)】]" and depth:
            depth -= 1
        if character in separators and depth == 0:
            value = re.sub(r"^[\s•·●▪■\-—]+|[\s。；;]+$", "", "".join(current)).strip()
            if value:
                values.append(value)
            current = []
        else:
            current.append(character)
    value = re.sub(r"^[\s•·●▪■\-—]+|[\s。；;]+$", "", "".join(current)).strip()
    if value:
        values.append(value)
    return values


def extract_recruitment_catalog(text: str) -> dict[str, list[str]]:
    """Extract explicit job and major lists from a recruitment source."""
    if not text:
        return {"job_titles": [], "major_requirements": []}
    job_section = _catalog_section(text, r"(?:招聘岗位|招聘职位|岗位需求|职位需求|招收岗位|招聘方向)")
    major_section = _catalog_section(text, r"(?:需求专业|专业需求|专业要求|招聘专业|专业类别|需求学科|所需专业|面向专业|专业方向|岗位专业)")
    job_titles: list[str] = []
    for line in job_section.splitlines():
        clean = re.sub(r"^\s*(?:[•·●▪■\-—]|\d+[.)、])\s*", "", line).strip()
        if not clean:
            continue
        job_titles.extend(_split_top_level(clean, "、；;•·，"))
    job_titles = list(dict.fromkeys(value for value in job_titles if 2 <= len(value) <= 120))

    major_requirements: list[str] = []
    for line in major_section.splitlines():
        clean = re.sub(r"^\s*(?:[•·●▪■\-—]|\d+[.)、])\s*", "", line).strip()
        if not clean or clean.replace("｜", "|") in {"专业类别|专业名称", "专业类别|专业方向"}:
            continue
        if re.fullmatch(r"(?:专业类别|专业名称|需求专业|专业要求)[:：]?", clean):
            continue
        major_requirements.append(clean.rstrip("。；;"))
    major_requirements = list(dict.fromkeys(value for value in major_requirements if 2 <= len(value) <= 500))
    return {"job_titles": job_titles, "major_requirements": major_requirements}


class TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.links: list[str] = []
        self.images: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript", "svg"}:
            self._skip += 1
        attributes = {key.lower(): value for key, value in attrs if value}
        if tag == "a" and attributes.get("href"):
            self.links.append(str(attributes["href"]))
        if tag == "img":
            for key in ("src", "data-src", "data-original", "data-lazy-src"):
                value = attributes.get(key)
                if value and not str(value).lower().startswith("data:"):
                    self.images.append(str(value))
                    break
            else:
                srcset = attributes.get("srcset") or ""
                if srcset:
                    value = str(srcset).split(",", 1)[0].strip().split(" ", 1)[0]
                    if value and not value.lower().startswith("data:"):
                        self.images.append(value)
        if tag == "meta":
            property_name = str(attributes.get("property") or attributes.get("name") or "").lower()
            content = attributes.get("content")
            if property_name in {"og:image", "twitter:image"} and content and not str(content).lower().startswith("data:"):
                self.images.append(str(content))

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript", "svg"} and self._skip:
            self._skip -= 1

    def handle_data(self, data: str) -> None:
        if not self._skip and data.strip():
            self.parts.append(data.strip())


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def detect_image_suffix(data: bytes) -> str | None:
    """Return a common image suffix when a response has no useful MIME type."""
    if data.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if data.startswith(b"\xff\xd8\xff"):
        return ".jpg"
    if data.startswith((b"GIF87a", b"GIF89a")):
        return ".gif"
    if data.startswith(b"BM"):
        return ".bmp"
    if len(data) >= 12 and data.startswith(b"RIFF") and data[8:12] == b"WEBP":
        return ".webp"
    if data.startswith((b"II*\x00", b"MM\x00*")):
        return ".tiff"
    return None


def extract_html(html: str) -> dict[str, Any]:
    parser = TextExtractor()
    parser.feed(html)
    text = "\n".join(parser.parts)
    title = parser.parts[0] if parser.parts else ""
    return {
        "title": title[:300],
        "text": text[:2_000_000],
        "links": list(dict.fromkeys(parser.links))[:100],
        "images": list(dict.fromkeys(parser.images))[:50],
    }


def _safe_decode(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030", "utf-16", "big5"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_file(filename: str, data: bytes) -> dict[str, Any]:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".log", ".csv", ".tsv"}:
        text = _safe_decode(data)
        if suffix in {".csv", ".tsv"}:
            delimiter = "\t" if suffix == ".tsv" else ","
            rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))[:20_000]
            text = "\n".join(" | ".join(row) for row in rows)
        return {"text": text[:2_000_000], "metadata": {"format": suffix[1:]}}
    if suffix == ".docx":
        try:
            from docx import Document

            document = Document(io.BytesIO(data))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            tables = []
            for table in document.tables:
                for row in table.rows:
                    tables.append(" | ".join(cell.text for cell in row.cells))
            return {"text": "\n".join(paragraphs + tables)[:2_000_000], "metadata": {"format": "docx"}}
        except Exception as exc:
            return {"text": "", "metadata": {"format": "docx", "error": str(exc)}}
    if suffix == ".xlsx":
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in workbook.worksheets[:20]:
                lines.append(f"[工作表: {sheet.title}]")
                for row in list(sheet.iter_rows(values_only=True))[:20_000]:
                    values = [str(value) for value in row if value is not None]
                    if values:
                        lines.append(" | ".join(values))
            return {"text": "\n".join(lines)[:2_000_000], "metadata": {"format": "xlsx"}}
        except Exception as exc:
            return {"text": "", "metadata": {"format": "xlsx", "error": str(exc)}}
    if suffix == ".pdf":
        try:
            import fitz

            document = fitz.open(stream=data, filetype="pdf")
            lines = []
            for page in list(document)[:200]:
                lines.append(page.get_text("text"))
            return {"text": "\n".join(lines)[:2_000_000], "metadata": {"format": "pdf", "pages": len(document)}}
        except Exception as exc:
            return {"text": "", "metadata": {"format": "pdf", "error": str(exc)}}
    if suffix in {".jpg", ".jpeg", ".png", ".webp", ".bmp"}:
        return process_image(data, suffix[1:])
    return {"text": "", "metadata": {"format": suffix.lstrip(".") or "unknown", "unsupported": True}}


def process_image(data: bytes, image_format: str) -> dict[str, Any]:
    """Run optional local OCR and QR decoding without making either mandatory."""
    try:
        import cv2
        import numpy as np

        image = cv2.imdecode(np.frombuffer(data, dtype=np.uint8), cv2.IMREAD_COLOR)
        if image is None:
            return {"text": "", "metadata": {"format": image_format, "decode_error": "invalid image"}}
        qr_values: list[str] = []
        detector = cv2.QRCodeDetector()
        try:
            ok, decoded, _, _ = detector.detectAndDecodeMulti(image)
            if ok:
                qr_values.extend(value for value in decoded if value)
            else:
                value, _, _ = detector.detectAndDecode(image)
                if value:
                    qr_values.append(value)
        except Exception:
            pass
        ocr_lines: list[str] = []
        try:
            from rapidocr_onnxruntime import RapidOCR

            result, _ = RapidOCR()(image)
            for item in result or []:
                if len(item) >= 2 and item[1]:
                    ocr_lines.append(str(item[1]))
        except Exception:
            pass
        return {
            "text": "\n".join(ocr_lines)[:2_000_000],
            "metadata": {"format": image_format, "needs_ocr": False},
            "qr_values": list(dict.fromkeys(qr_values)),
        }
    except Exception as exc:
        return {"text": "", "metadata": {"format": image_format, "ocr_error": str(exc)}, "qr_values": []}


def validate_public_url(url: str) -> str:
    parsed_url = urlparse(url)
    if parsed_url.scheme not in {"http", "https"} or not parsed_url.hostname:
        raise ValueError("Only public HTTP(S) URLs are supported")
    if parsed_url.username or parsed_url.password:
        raise ValueError("URLs with embedded credentials are not supported")
    host = parsed_url.hostname.lower()
    if host in {"localhost", "metadata.google.internal"}:
        raise ValueError("Local and metadata hosts are blocked")
    try:
        port = parsed_url.port or (443 if parsed_url.scheme == "https" else 80)
    except ValueError as exc:
        raise ValueError("Invalid URL port") from exc
    try:
        address = ipaddress.ip_address(host)
        if address.is_private or address.is_loopback or address.is_link_local or address.is_reserved:
            raise ValueError("Private network targets are blocked")
    except ValueError as exc:
        if str(exc) == "Private network targets are blocked":
            raise
    try:
        addresses = {info[4][0] for info in socket.getaddrinfo(host, port, type=socket.SOCK_STREAM)}
        if not addresses:
            raise ValueError("Unable to resolve public URL")
        for resolved in addresses:
            address = ipaddress.ip_address(resolved)
            if address.is_private or address.is_loopback or address.is_link_local or address.is_reserved:
                raise ValueError("URL resolves to a private network target")
    except socket.gaierror as exc:
        raise ValueError("Unable to resolve public URL") from exc
    return url


def recover_original_source_url(value: Any) -> str | None:
    """Recover the original WeChat URL from a temporary verification redirect."""
    candidate = str(value or "").strip()
    if not candidate:
        return None
    for _ in range(2):
        parsed = urlparse(candidate)
        host = (parsed.hostname or "").lower()
        if host != "mp.weixin.qq.com" and not host.endswith(".weixin.qq.com"):
            break
        target_values = parse_qs(parsed.query).get("target_url")
        if not target_values:
            break
        target = str(target_values[0] or "").strip()
        target_parsed = urlparse(target)
        if not target or target == candidate or target_parsed.scheme not in {"http", "https"} or not target_parsed.hostname:
            break
        candidate = target
    return candidate


def fetch_public_http(url: str, timeout: float = 30, max_bytes: int = 10 * 1024 * 1024) -> httpx.Response:
    current_url = url
    for _ in range(6):
        validate_public_url(current_url)
        response = httpx.get(
            current_url,
            timeout=timeout,
            follow_redirects=False,
            headers={"User-Agent": "JobPostings/0.1 (+local recruitment archive)"},
        )
        if 300 <= response.status_code < 400:
            location = response.headers.get("location")
            if not location:
                raise ValueError("Redirect response has no location")
            current_url = urljoin(current_url, location)
            continue
        response.raise_for_status()
        if len(response.content) > max_bytes:
            raise ValueError(f"Response is larger than {max_bytes // (1024 * 1024)} MB")
        return response
    raise ValueError("Too many redirects")


def is_access_challenge_page(url: str, text: str) -> bool:
    """Detect the WeChat environment-verification page instead of treating it as article text."""
    host = (urlparse(url).hostname or "").lower()
    if host != "mp.weixin.qq.com" and not host.endswith(".weixin.qq.com"):
        return False
    compact = normalize_text(text)
    return "当前环境异常" in compact and "完成验证后即可继续访问" in compact


def fetch_public_url(url: str) -> dict[str, Any]:
    response = fetch_public_http(url)
    content_type = response.headers.get("content-type", "")
    media_type = content_type.split(";", 1)[0].strip().lower()
    detected_suffix = detect_image_suffix(response.content)
    if media_type.startswith("image/") or detected_suffix:
        suffix = mimetypes.guess_extension(media_type) or detected_suffix or Path(urlparse(str(response.url)).path).suffix or ".bin"
        return {
            "url": str(response.url),
            "text": "",
            "content_type": content_type,
            "images": [],
            "data": response.content,
            "filename": f"web-image{suffix}",
        }
    if is_access_challenge_page(str(response.url), response.text):
        return {
            "url": str(response.url),
            "text": "",
            "content_type": content_type,
            "images": [],
            "access_challenge": True,
            "access_error": "微信返回环境验证页面",
        }
    if "html" not in content_type and not response.text.lstrip().startswith("<"):
        return {"url": str(response.url), "text": response.text[:2_000_000], "content_type": content_type}
    result = extract_html(response.text)
    result["links"] = [urljoin(str(response.url), link) for link in result.get("links", []) if link]
    result["images"] = [urljoin(str(response.url), image) for image in result.get("images", []) if image]
    result.update({"url": str(response.url), "content_type": content_type})
    return result


def _mapping(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if not isinstance(value, str) or len(value) > 2_000_000:
        return {}
    candidate = value.strip()
    if not candidate:
        return {}
    for loader in (json.loads, ast.literal_eval):
        try:
            result = loader(candidate)
        except (ValueError, SyntaxError, TypeError, json.JSONDecodeError):
            continue
        if isinstance(result, dict):
            return result
    return {}


def content_data(message: dict[str, Any]) -> dict[str, Any]:
    return _mapping(message.get("contentData") or message.get("content_data"))


def normalized_message_type(value: Any) -> str:
    return str(value or "text").strip().lower()


SYSTEM_MESSAGE_TYPES = frozenset({
    "system",
    "system_message",
    "system message",
    "system_notification",
    "system-notification",
    "notification",
    "group_notification",
    "group_notice",
    "sysmsg",
    "sys_msg",
    "系统",
    "系统消息",
    "系统信息",
    "系统通知",
    "通知",
    "群通知",
    "群消息通知",
    "群聊通知",
    "群聊系统通知",
    "群系统消息",
})
LINK_MESSAGE_TYPES = frozenset({
    "article",
    "link",
    "url",
    "公众号链接",
    "分享消息",
    "分享",
})
IMAGE_MESSAGE_TYPES = frozenset({
    "image",
    "picture",
    "photo",
    "图片",
    "图像",
})
FILE_MESSAGE_TYPES = frozenset({
    "file",
    "attachment",
    "document",
    "文件",
    "文档",
})


def is_system_message(message_type: Any, text: str = "", message: dict[str, Any] | None = None) -> bool:
    normalized = normalized_message_type(message_type)
    if normalized in SYSTEM_MESSAGE_TYPES:
        return True
    message = message or {}
    if str(message.get("from") or message.get("source") or "").strip().lower() == "system":
        return True
    nested = content_data(message)
    if str(nested.get("type") or "").strip().lower() in SYSTEM_MESSAGE_TYPES:
        return True
    compact = normalize_text(text)
    return bool(re.search(
        r"(?:邀请.{0,120}(?:加入|进入)群聊|(?:加入|进入)了群聊|(?:退出|离开)了群聊|(?:被|将|把).{0,100}移出群聊|撤回了?(?:一条)?消息|修改群名|设置了群公告|拍了拍)",
        compact,
    ))


def is_link_message(message_type: Any, metadata: dict[str, Any] | None = None) -> bool:
    normalized = normalized_message_type(message_type)
    if normalized in LINK_MESSAGE_TYPES:
        return True
    if metadata and metadata.get("url"):
        return True
    nested = _mapping((metadata or {}).get("contentData") or (metadata or {}).get("content_data"))
    return str(nested.get("type") or "").strip().lower() in {"share", "link", "url", "article"}


def is_image_message(message_type: Any, message: dict[str, Any] | None = None) -> bool:
    normalized = normalized_message_type(message_type)
    if normalized in IMAGE_MESSAGE_TYPES:
        return True
    message = message or {}
    nested = content_data(message)
    media = _mapping(message.get("media"))
    return str(nested.get("type") or "").strip().lower() == "image" or str(media.get("type") or "").strip().lower() == "image"


def is_file_message(message_type: Any, message: dict[str, Any] | None = None) -> bool:
    normalized = normalized_message_type(message_type)
    if normalized in FILE_MESSAGE_TYPES:
        return True
    nested = content_data(message or {})
    return str(nested.get("type") or "").strip().lower() in {"file", "document", "attachment"}


def is_text_message(message_type: Any, message: dict[str, Any] | None = None) -> bool:
    return not (
        is_link_message(message_type, message)
        or is_image_message(message_type, message)
        or is_file_message(message_type, message)
    )


def parse_message_payload(message: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    message_type = normalized_message_type(message.get("type") or message.get("msgType") or "text")
    text = str(message.get("text") or message.get("content") or "")
    metadata = dict(message)
    metadata.pop("text", None)
    metadata.pop("content", None)
    nested = content_data(message)
    if nested:
        metadata["contentData"] = nested
    url = message.get("url") or nested.get("url")
    if url:
        metadata["url"] = str(url)
        metadata["source_url"] = str(url)
    if nested.get("title"):
        metadata["shared_title"] = str(nested["title"])
    if nested.get("des") or nested.get("description"):
        metadata["shared_description"] = str(nested.get("des") or nested.get("description"))
    segments = [normalize_text(text)]
    if is_link_message(message_type, metadata):
        segments.extend(normalize_text(str(value)) for value in (nested.get("title"), nested.get("des") or nested.get("description")) if value)
    deduplicated = list(dict.fromkeys(value for value in segments if value))
    return "\n".join(deduplicated), metadata
